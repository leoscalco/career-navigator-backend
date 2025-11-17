"""
Document parser for extracting text from PDF and Word documents.
"""

import io
from typing import Optional
from pathlib import Path


class DocumentParser:
    """Parser for extracting text from various document formats."""
    
    @staticmethod
    def parse_pdf(file_content: bytes) -> str:
        """
        Extract text from PDF file.
        
        Args:
            file_content: PDF file bytes
            
        Returns:
            Extracted text content
        """
        try:
            import pypdf
            
            pdf_file = io.BytesIO(file_content)
            pdf_reader = pypdf.PdfReader(pdf_file)
            
            text_parts = []
            for page in pdf_reader.pages:
                text = page.extract_text()
                if text:
                    text_parts.append(text)
            
            return "\n\n".join(text_parts)
        except ImportError:
            # Fallback to pdfplumber if pypdf fails
            try:
                import pdfplumber
                
                pdf_file = io.BytesIO(file_content)
                text_parts = []
                
                with pdfplumber.open(pdf_file) as pdf:
                    for page in pdf.pages:
                        text = page.extract_text()
                        if text:
                            text_parts.append(text)
                
                return "\n\n".join(text_parts)
            except ImportError:
                raise ImportError(
                    "PDF parsing requires either 'pypdf' or 'pdfplumber'. "
                    "Install with: poetry add pypdf pdfplumber"
                )
        except Exception as e:
            raise ValueError(f"Failed to parse PDF: {str(e)}")
    
    @staticmethod
    def parse_docx(file_content: bytes) -> str:
        """
        Extract text from Word (.docx) file.
        
        Args:
            file_content: DOCX file bytes
            
        Returns:
            Extracted text content
        """
        try:
            from docx import Document
            
            docx_file = io.BytesIO(file_content)
            doc = Document(docx_file)
            
            text_parts = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_parts.append(" | ".join(row_text))
            
            return "\n\n".join(text_parts)
        except ImportError:
            raise ImportError(
                "Word document parsing requires 'python-docx'. "
                "Install with: poetry add python-docx"
            )
        except Exception as e:
            raise ValueError(f"Failed to parse Word document: {str(e)}")
    
    @staticmethod
    def parse_document(file_content: bytes, filename: str) -> str:
        """
        Parse document based on file extension.
        
        Args:
            file_content: Document file bytes
            filename: Original filename (used to determine format)
            
        Returns:
            Extracted text content
        """
        file_path = Path(filename.lower())
        extension = file_path.suffix.lower()
        
        if extension == ".pdf":
            return DocumentParser.parse_pdf(file_content)
        elif extension in [".docx", ".doc"]:
            if extension == ".doc":
                raise ValueError(
                    "Old .doc format is not supported. Please convert to .docx format."
                )
            return DocumentParser.parse_docx(file_content)
        elif extension in [".txt", ".text"]:
            # Plain text file
            try:
                return file_content.decode("utf-8")
            except UnicodeDecodeError:
                # Try other encodings
                for encoding in ["latin-1", "cp1252", "iso-8859-1"]:
                    try:
                        return file_content.decode(encoding)
                    except UnicodeDecodeError:
                        continue
                raise ValueError("Could not decode text file. Please use UTF-8 encoding.")
        else:
            raise ValueError(
                f"Unsupported file format: {extension}. "
                "Supported formats: .pdf, .docx, .txt"
            )
    
    @staticmethod
    def get_supported_formats() -> list[str]:
        """Get list of supported file formats."""
        return [".pdf", ".docx", ".txt"]

